"""
Write a position paper string to a .docx file with proper MUN formatting.
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, name: str = "Times New Roman", size: int = 12):
    run.font.name = name
    run.font.size = Pt(size)
    # Force the theme font so Word doesn't override with Calibri
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _fmt(para, first_line: float = 0.0, left: float = 0.0, hanging: float = 0.0):
    """Apply spacing and indent via python-docx paragraph_format (reliable)."""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if hanging:
        pf.left_indent = Inches(left)
        pf.first_line_indent = -Inches(hanging)
    elif first_line:
        pf.first_line_indent = Inches(first_line)
    else:
        pf.first_line_indent = Inches(0)


SECTION_HEADER_RE = re.compile(r"^\d+\.\s+\S")
WORKS_CITED_RE = re.compile(r"^works cited", re.IGNORECASE)


def write_docx(paper_text: str, output_path: str) -> None:
    doc = Document()

    # 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Override Normal style defaults so spacing is clean everywhere
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    in_works_cited = False

    for raw_line in paper_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        para = doc.add_paragraph()

        if WORKS_CITED_RE.match(line):
            in_works_cited = True
            run = para.add_run(line)
            run.bold = True
            _set_font(run)
            _fmt(para)
            continue

        if in_works_cited:
            run = para.add_run(line)
            _set_font(run)
            _fmt(para, left=0.5, hanging=0.5)
            continue

        if SECTION_HEADER_RE.match(line):
            run = para.add_run(line)
            run.bold = True
            _set_font(run)
            _fmt(para)
            continue

        # Body paragraph — 0.5in first-line indent
        run = para.add_run(line)
        _set_font(run)
        _fmt(para, first_line=0.5)

    doc.save(output_path)
